import os
import json
import time
import re
import requests
import markdown
from io import BytesIO
from flask import Flask, request, jsonify, render_template, send_file, Response, stream_with_context
from dotenv import load_dotenv
from groq import Groq
from tavily import TavilyClient
import PyPDF2
import docx

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-key")

client = Groq(api_key=os.getenv("GROQ_API_KEY"))
SERPER_API_KEY = os.getenv("SERPER_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
SEARCH_PROVIDER = os.getenv("SEARCH_PROVIDER", "").lower()  # 'serper', 'tavily', 'both', or '' (auto-detect)

tavily_client = TavilyClient(api_key=TAVILY_API_KEY) if TAVILY_API_KEY else None

# Fast, stable model — no compound/web-search model, we handle search ourselves
MODEL = "llama-3.3-70b-versatile"

MODES = {
    "deep_research": {
        "name": "Deep Research",
        "description": "Comprehensive multi-angle research report on any topic",
        "icon": "radar",
        "search_queries": 4,
        "system": """You are an elite research analyst and academic writer. You will be given a topic along with fresh web search results as context. Your task is to produce an extraordinarily comprehensive, well-structured research report.

For every topic:
1. **EXECUTIVE SUMMARY** — 3-5 sentence overview
2. **HISTORICAL CONTEXT & BACKGROUND** — origins, timeline, key milestones
3. **CURRENT STATE & KEY DEVELOPMENTS** — what's happening right now, use the provided search context heavily
4. **CORE CONCEPTS & MECHANISMS** — how it actually works
5. **MULTIPLE PERSPECTIVES & DEBATES** — contrasting views, controversies
6. **REAL-WORLD APPLICATIONS & CASE STUDIES** — concrete examples
7. **CHALLENGES & LIMITATIONS** — what's broken or contested
8. **FUTURE OUTLOOK & EMERGING TRENDS** — where it's heading
9. **CONCLUSIONS & KEY TAKEAWAYS** — crisp synthesis

Use markdown with proper headers (##, ###), bullet points, bold terms, tables where relevant, and code blocks where applicable. Be exhaustive, authoritative, and cite specific examples, dates, names, and data points. Use the search results to include current, accurate information."""
    },
    "paper_crux": {
        "name": "Paper Crux",
        "description": "Extract essential insights from any research paper",
        "icon": "flask",
        "search_queries": 2,
        "system": """You are an expert academic synthesizer. You will be given a research paper or academic content, plus optional web context about the topic. Make complex research instantly understandable.

1. **THE CORE QUESTION** — What fundamental problem does this paper address?
2. **WHY IT MATTERS** — Real-world significance and why this research was needed
3. **METHODOLOGY IN PLAIN ENGLISH** — How they did it, explained simply
4. **KEY FINDINGS** — The most important discoveries, in order of significance
5. **THE BREAKTHROUGH** (if any) — What's genuinely new here?
6. **LIMITATIONS & CAVEATS** — What the paper doesn't cover or gets wrong
7. **PRACTICAL IMPLICATIONS** — How can this be applied or used?
8. **VERDICT** — Is this paper's contribution significant? Why?
9. **5-SENTENCE SUMMARY** — The entire paper distilled to 5 sentences

Use markdown. Make complex jargon instantly clear. Be ruthlessly insightful."""
    },
    "docs_simplifier": {
        "name": "Docs Simplifier",
        "description": "Transform complex documentation into crisp, actionable guides",
        "icon": "book",
        "search_queries": 2,
        "system": """You are a master technical writer. You will be given documentation or technical content, plus optional web context. Transform it into a crystal-clear guide.

1. **WHAT THIS IS** — One-line explanation
2. **QUICK START** — The fastest path to getting it working
3. **CORE CONCEPTS** — Only the concepts you MUST understand, explained cleanly
4. **KEY COMMANDS/FUNCTIONS/METHODS** — A clean reference table or list
5. **COMMON PATTERNS** — The 20% of features you'll use 80% of the time
6. **GOTCHAS & PITFALLS** — What will trip you up (save hours of debugging)
7. **PRACTICAL EXAMPLES** — Real, runnable examples for each major concept
8. **CHEAT SHEET** — Ultra-condensed reference at the end

Use markdown with code blocks. Be precise, not wordy. Every sentence must earn its place."""
    },
    "custom": {
        "name": "Custom Analysis",
        "description": "Custom prompt-based deep analysis with web context",
        "icon": "target",
        "search_queries": 3,
        "system": """You are an elite analyst capable of any form of deep research and synthesis. You will be given a topic, optional web search results, and a custom instruction defining your analytical lens. Follow the custom instruction precisely while producing a thorough, well-structured, markdown-formatted response. Use headers, bullet points, tables, and code blocks as appropriate. Be comprehensive, precise, and genuinely insightful."""
    }
}


# ─── Serper Search ────────────────────────────────────────────────────────────

def serper_search(query, num_results=8):
    """Fetch web search results from Serper API."""
    if not SERPER_API_KEY:
        return []
    try:
        response = requests.post(
            "https://google.serper.dev/search",
            headers={
                "X-API-KEY": SERPER_API_KEY,
                "Content-Type": "application/json"
            },
            json={"q": query, "num": num_results},
            timeout=10
        )
        response.raise_for_status()
        data = response.json()
        results = []

        # Knowledge graph
        if "knowledgeGraph" in data:
            kg = data["knowledgeGraph"]
            if kg.get("description"):
                results.append({
                    "title": kg.get("title", "Knowledge Graph"),
                    "snippet": kg.get("description", ""),
                    "url": kg.get("website", ""),
                    "type": "knowledge_graph"
                })

        # Organic results
        for r in data.get("organic", []):
            results.append({
                "title": r.get("title", ""),
                "snippet": r.get("snippet", ""),
                "url": r.get("link", ""),
                "date": r.get("date", ""),
                "type": "organic"
            })

        # Answer box
        if "answerBox" in data:
            ab = data["answerBox"]
            snippet = ab.get("snippet") or ab.get("answer") or ab.get("snippetHighlighted", [""])[0]
            if snippet:
                results.insert(0, {
                    "title": ab.get("title", "Direct Answer"),
                    "snippet": snippet if isinstance(snippet, str) else " ".join(snippet),
                    "url": ab.get("link", ""),
                    "type": "answer_box"
                })

        return results[:num_results]
    except Exception as e:
        print(f"[Serper Error] {e}")
        return []


# ─── Tavily Search ───────────────────────────────────────────────────────────

def tavily_search(query, num_results=8):
    """Fetch web search results from Tavily API, normalised to the same schema as serper_search()."""
    if not tavily_client:
        return []
    try:
        response = tavily_client.search(
            query=query,
            max_results=num_results,
            search_depth="advanced",
        )
        results = []

        for r in response.get("results", []):
            results.append({
                "title": r.get("title", ""),
                "snippet": r.get("content", ""),
                "url": r.get("url", ""),
                "date": r.get("published_date", ""),
                "type": "organic"
            })

        return results[:num_results]
    except Exception as e:
        print(f"[Tavily Error] {e}")
        return []


# ─── Search Dispatcher ───────────────────────────────────────────────────────

def web_search(query, num_results=8):
    """Route search to serper, tavily, or both based on config and available keys."""
    provider = SEARCH_PROVIDER

    # Auto-detect when no explicit provider is set
    if not provider:
        if TAVILY_API_KEY and SERPER_API_KEY:
            provider = "both"
        elif TAVILY_API_KEY:
            provider = "tavily"
        elif SERPER_API_KEY:
            provider = "serper"
        else:
            return []

    if provider == "tavily":
        return tavily_search(query, num_results)
    elif provider == "serper":
        return serper_search(query, num_results)
    elif provider == "both":
        serper_results = serper_search(query, num_results)
        tavily_results = tavily_search(query, num_results)
        # Merge, dedup by URL, prefer serper ordering first
        seen_urls = set()
        merged = []
        for r in serper_results + tavily_results:
            url = r.get("url", "")
            key = url if url else r.get("snippet", "")[:80]
            if key not in seen_urls:
                seen_urls.add(key)
                merged.append(r)
        return merged[:num_results]
    else:
        # Unknown provider value, fall back to serper
        return serper_search(query, num_results)


def build_search_queries(mode, topic, custom_instruction=None):
    """Generate smart search queries based on mode and topic."""
    num = MODES.get(mode, MODES["deep_research"])["search_queries"]

    if mode == "deep_research":
        return [
            topic,
            f"{topic} latest developments 2024 2025",
            f"{topic} research analysis statistics",
            f"{topic} challenges future trends"
        ][:num]
    elif mode == "paper_crux":
        return [
            topic,
            f"{topic} research findings"
        ][:num]
    elif mode == "docs_simplifier":
        return [
            f"{topic} documentation guide",
            f"{topic} examples tutorial"
        ][:num]
    elif mode == "custom" and custom_instruction:
        return [
            topic,
            f"{topic} {custom_instruction[:60]}",
            f"{topic} analysis data"
        ][:num]
    else:
        return [topic]


def format_search_context(results_by_query):
    """Format all search results into a clean context block for the LLM."""
    if not any(results_by_query.values()):
        return ""

    lines = ["## WEB SEARCH CONTEXT (use this for current, accurate information)\n"]

    seen_urls = set()
    result_num = 1

    for query, results in results_by_query.items():
        if results:
            lines.append(f"### Search: \"{query}\"\n")
            for r in results:
                url = r.get("url", "")
                if url in seen_urls:
                    continue
                seen_urls.add(url)

                title = r.get("title", "")
                snippet = r.get("snippet", "")
                date = r.get("date", "")
                r_type = r.get("type", "organic")

                if r_type == "answer_box":
                    lines.append(f"**[DIRECT ANSWER]** {title}")
                elif r_type == "knowledge_graph":
                    lines.append(f"**[KNOWLEDGE PANEL]** {title}")
                else:
                    lines.append(f"**[{result_num}]** {title}")

                if date:
                    lines.append(f"*Published: {date}*")
                lines.append(snippet)
                if url:
                    lines.append(f"Source: {url}")
                lines.append("")
                result_num += 1

    lines.append("---\nEnd of web context. Now write your report based on the above information and your knowledge.\n")
    return "\n".join(lines)


# ─── File Extraction ──────────────────────────────────────────────────────────

def extract_text_from_file(file):
    filename = file.filename.lower()
    content = ""

    if filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        except Exception as e:
            content = f"[PDF extraction error: {str(e)}]"
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        try:
            doc = docx.Document(file)
            for para in doc.paragraphs:
                content += para.text + "\n"
        except Exception as e:
            content = f"[DOCX extraction error: {str(e)}]"
    elif filename.endswith('.md'):
        content = file.read().decode('utf-8', errors='ignore')
    elif filename.endswith('.txt'):
        content = file.read().decode('utf-8', errors='ignore')
    else:
        try:
            content = file.read().decode('utf-8', errors='ignore')
        except Exception:
            content = "[Unable to read file]"

    return content.strip()


def build_prompt(mode, topic, file_content=None, custom_instruction=None, search_context=""):
    parts = []

    if search_context:
        parts.append(search_context)

    if file_content:
        if mode == "paper_crux":
            parts.append(f"## DOCUMENT TO ANALYZE\n\n{file_content[:12000]}")
        elif mode == "docs_simplifier":
            parts.append(f"## DOCUMENTATION TO SIMPLIFY\n\n{file_content[:12000]}")
        else:
            parts.append(f"## UPLOADED DOCUMENT CONTENT\n\n{file_content[:8000]}")

    if mode == "custom" and custom_instruction:
        parts.append(f"## YOUR ANALYTICAL INSTRUCTION\n\n{custom_instruction}")

    parts.append(f"## TOPIC / QUERY\n\n{topic}")
    parts.append("Now produce your complete, thorough response:")

    return "\n\n".join(parts)


# ─── MinimalPDF ───────────────────────────────────────────────────────────────

class MinimalPDF(object):
    PW = 595
    PH = 842
    ML = 50
    MR = 50
    MT = 50
    MB = 50

    def __init__(self):
        self._objects = []
        self._offsets = []
        self._pages = []
        self._buf = []
        self._y = 0
        self._font = None
        self._font_size = 0

    def _add_obj(self, content):
        self._objects.append(content)
        return len(self._objects)

    def _w(self):
        return self.PW - self.ML - self.MR

    @staticmethod
    def _pdf_str(s):
        s = s.encode('latin-1', 'replace').decode('latin-1')
        s = s.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
        return '(' + s + ')'

    @staticmethod
    def _char_width(ch, font_size, bold=False):
        wide = 'WMmwABCDEFGHIJKLNOPQRSTUVXYZ'
        narrow = 'iIl1!|;:,. '
        if ch in wide:
            ratio = 0.72
        elif ch in narrow:
            ratio = 0.32
        else:
            ratio = 0.55
        return font_size * ratio

    def _text_width(self, text, font_size, bold=False):
        return sum(self._char_width(c, font_size, bold) for c in text)

    def _wrap_text(self, text, font_size, bold=False, indent=0):
        max_w = self._w() - indent
        words = text.split(' ')
        lines = []
        current = ''
        for word in words:
            test = (current + ' ' + word).strip()
            if self._text_width(test, font_size, bold) <= max_w:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines if lines else ['']

    def _new_page(self):
        if self._buf:
            self._flush_page()
        self._buf = []
        self._y = self.MT

    def _flush_page(self):
        stream = '\n'.join(self._buf)
        stream_bytes = stream.encode('latin-1', 'replace')
        cs_idx = self._add_obj(
            '<< /Length {} >>\nstream\n'.format(len(stream_bytes)).encode('latin-1')
            + stream_bytes
            + b'\nendstream'
        )
        page_idx = self._add_obj(
            '<< /Type /Page /Parent 2 0 R'
            ' /MediaBox [0 0 {} {}]'
            ' /Contents {} 0 R'
            ' /Resources << /Font << /F1 4 0 R /F2 5 0 R /F3 6 0 R /F4 7 0 R >> >>'
            ' >>'.format(self.PW, self.PH, cs_idx)
        )
        self._pages.append(page_idx)
        self._buf = []

    def _check_page(self, needed):
        if self._y + needed > self.PH - self.MB:
            self._flush_page()
            self._y = self.MT

    def _draw_line(self, text, x, y, font, size, r=0, g=0, b=0):
        font_map = {'H': '/F1', 'HB': '/F2', 'HI': '/F3', 'C': '/F4'}
        pdf_y = self.PH - y
        self._buf.append(
            '{:.3f} {:.3f} {:.3f} rg BT {} {} Tf {:.2f} {:.2f} Td {} Tj ET 0 0 0 rg'.format(
                r, g, b, font_map[font], size, x, pdf_y,
                self._pdf_str(text)
            )
        )

    def _draw_rect_filled(self, x, y, w, h, r=0, g=0, b=0):
        pdf_y = self.PH - y - h
        self._buf.append(
            '{:.3f} {:.3f} {:.3f} rg {:.2f} {:.2f} {:.2f} {:.2f} re f 0 0 0 rg'.format(
                r, g, b, x, pdf_y, w, h
            )
        )

    def _draw_hline(self, y, r=0.7, g=0.7, b=0.7):
        pdf_y = self.PH - y
        self._buf.append(
            '{:.3f} {:.3f} {:.3f} RG 0.5 w {} {} m {} {} l S 0 0 0 RG'.format(
                r, g, b,
                self.ML, pdf_y, self.PW - self.MR, pdf_y
            )
        )

    def add_title_block(self, title):
        self._check_page(50)
        self._draw_rect_filled(0, self._y, self.PW, 44, 0.07, 0.07, 0.12)
        tw = self._text_width(title[:70], 18, bold=True)
        tx = max(self.ML, (self.PW - tw) / 2)
        self._draw_line(title[:70], tx, self._y + 14, 'HB', 18, 1, 1, 1)
        self._y += 52

    def add_heading(self, text, level):
        sizes = {
            1: (16, 'HB', 0.07, 0.07, 0.15),
            2: (13, 'HB', 0.10, 0.31, 0.63),
            3: (11, 'HB', 0.18, 0.42, 0.75),
            4: (10, 'HB', 0.33, 0.33, 0.33)
        }
        size, font, r, g, b = sizes.get(level, sizes[4])
        before = {1: 12, 2: 9, 3: 6, 4: 4}[level]
        self._y += before
        lines = self._wrap_text(text, size)
        for ln in lines:
            self._check_page(size + 6)
            self._draw_line(ln, self.ML, self._y, font, size, r, g, b)
            self._y += size + 4
        if level == 1:
            self._y += 2
            self._draw_hline(self._y)
            self._y += 5

    def add_paragraph(self, text, indent=0, font='H', size=10,
                      r=0.13, g=0.13, b=0.13, line_gap=5):
        lines = self._wrap_text(text, size, indent=indent)
        for ln in lines:
            self._check_page(size + line_gap)
            self._draw_line(ln, self.ML + indent, self._y, font, size, r, g, b)
            self._y += size + line_gap
        self._y += 2

    def add_bullet(self, text):
        self._check_page(16)
        self._draw_line(u'\u2022', self.ML + 8, self._y, 'H', 10, 0.10, 0.31, 0.63)
        lines = self._wrap_text(text, 10, indent=22)
        for ln in lines:
            self._check_page(15)
            self._draw_line(ln, self.ML + 22, self._y, 'H', 10, 0.13, 0.13, 0.13)
            self._y += 15

    def add_code_block(self, text):
        lines = text.split('\n')
        total_h = len(lines) * 13 + 12
        if self._y + total_h < self.PH - self.MB:
            self._draw_rect_filled(self.ML, self._y, self._w(), total_h, 0.93, 0.93, 0.93)
            for ln in lines:
                self._draw_line(ln[:90], self.ML + 6, self._y + 4, 'C', 8, 0.10, 0.10, 0.10)
                self._y += 13
            self._y += 6
        else:
            for ln in lines:
                self._check_page(14)
                self._draw_line(ln[:90], self.ML + 6, self._y, 'C', 8, 0.10, 0.10, 0.10)
                self._y += 13
            self._y += 4

    def add_spacer(self, h=6):
        self._y += h

    def add_hr(self):
        self._y += 4
        self._draw_hline(self._y)
        self._y += 6

    def output(self):
        if self._buf or not self._pages:
            self._flush_page()

        out = BytesIO()
        out.write(b'%PDF-1.4\n')
        xref_offsets = {}

        def write_obj(num, content):
            xref_offsets[num] = out.tell()
            out.write('{} 0 obj\n'.format(num).encode())
            if isinstance(content, bytes):
                out.write(content)
            else:
                out.write(content.encode('latin-1', 'replace'))
            out.write(b'\nendobj\n')

        content_obj_start = 8
        shift = content_obj_start - 1

        def shifted(content):
            if isinstance(content, bytes):
                def replace_ref(m):
                    old_num = int(m.group(1))
                    return '{} 0 R'.format(old_num + shift)
                text = content.decode('latin-1')
                text = re.sub(r'(\d+) 0 R', replace_ref, text)
                return text.encode('latin-1')
            else:
                def replace_ref(m):
                    old_num = int(m.group(1))
                    return '{} 0 R'.format(old_num + shift)
                return re.sub(r'(\d+) 0 R', replace_ref, content)

        page_obj_numbers = [p + shift for p in self._pages]
        pages_kids = ' '.join('{} 0 R'.format(n) for n in page_obj_numbers)
        total_objs = 7 + len(self._objects)

        write_obj(1, '<< /Type /Catalog /Pages 2 0 R >>')
        write_obj(2, '<< /Type /Pages /Kids [{}] /Count {} >>'.format(pages_kids, len(self._pages)))
        write_obj(3, '<< /Producer (MAX-Research) >>')

        fonts = [
            (4, '/F1', 'Helvetica'),
            (5, '/F2', 'Helvetica-Bold'),
            (6, '/F3', 'Helvetica-Oblique'),
            (7, '/F4', 'Courier'),
        ]
        for num, alias, fname in fonts:
            write_obj(num, '<< /Type /Font /Subtype /Type1 /BaseFont /{} /Encoding /WinAnsiEncoding >>'.format(fname))

        for idx, obj_content in enumerate(self._objects):
            obj_num = content_obj_start + idx
            write_obj(obj_num, shifted(obj_content))

        xref_pos = out.tell()
        out.write('xref\n0 {}\n'.format(total_objs + 1).encode())
        out.write(b'0000000000 65535 f \n')
        for i in range(1, total_objs + 1):
            offset = xref_offsets.get(i, 0)
            out.write('{:010d} 00000 n \n'.format(offset).encode())

        out.write('trailer\n<< /Size {} /Root 1 0 R /Info 3 0 R >>\n'.format(total_objs + 1).encode())
        out.write('startxref\n{}\n%%EOF\n'.format(xref_pos).encode())
        out.seek(0)
        return out.read()


# ─── Export Helpers ───────────────────────────────────────────────────────────

def _strip_md(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text.strip()


def generate_pdf_from_markdown(md_text, title="Research Report"):
    pdf = MinimalPDF()
    pdf._new_page()
    pdf.add_title_block(title[:80])
    pdf.add_spacer(8)

    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                pdf.add_code_block('\n'.join(code_buf))
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped or stripped == '---':
            if stripped == '---':
                pdf.add_hr()
            else:
                pdf.add_spacer(4)
        elif stripped.startswith('#### '):
            pdf.add_heading(_strip_md(stripped[5:]), 4)
        elif stripped.startswith('### '):
            pdf.add_heading(_strip_md(stripped[4:]), 3)
        elif stripped.startswith('## '):
            pdf.add_heading(_strip_md(stripped[3:]), 2)
        elif stripped.startswith('# '):
            pdf.add_heading(_strip_md(stripped[2:]), 1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            pdf.add_bullet(_strip_md(stripped[2:]))
        elif re.match(r'^\d+\. ', stripped):
            pdf.add_bullet(_strip_md(re.sub(r'^\d+\. ', '', stripped)))
        elif stripped.startswith('> '):
            pdf.add_paragraph(_strip_md(stripped[2:]), indent=16, font='HI', r=0.33, g=0.33, b=0.33)
        else:
            pdf.add_paragraph(_strip_md(stripped))
        i += 1

    return pdf.output()


def generate_docx_from_markdown(md_text, title="Research Report"):
    doc = docx.Document()
    title_para = doc.add_heading(title, 0)
    title_para.alignment = 1

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_para = doc.add_paragraph('\n'.join(code_lines))
            code_para.style = 'No Spacing'
            run = code_para.runs[0] if code_para.runs else code_para.add_run('')
            run.font.name = 'Courier New'
            run.font.size = docx.shared.Pt(9)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, line[2:])
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, re.sub(r'^\d+\. ', '', line))
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        elif line.strip() == '' or line.strip() == '---':
            doc.add_paragraph('')
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_run(p, line)
        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_formatted_run(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html', modes=MODES)


@app.route('/api/modes')
def get_modes():
    return jsonify(MODES)


@app.route('/api/search-status')
def search_status():
    """Let the frontend know which search providers are configured."""
    return jsonify({
        "serper_enabled": bool(SERPER_API_KEY),
        "tavily_enabled": bool(TAVILY_API_KEY),
        "search_provider": SEARCH_PROVIDER or "auto"
    })


@app.route('/api/research/stream', methods=['POST'])
def research_stream():
    mode = request.form.get('mode', 'deep_research')
    topic = request.form.get('topic', '').strip()
    custom_instruction = request.form.get('custom_instruction', '').strip()

    file_content = None
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename:
            file_content = extract_text_from_file(file)

    if not topic and not file_content:
        return jsonify({"error": "Please provide a topic or upload a file"}), 400

    if not topic and file_content:
        topic = "Analyze the uploaded document"

    system_prompt = MODES.get(mode, MODES['deep_research'])['system']

    def generate():
        try:
            # ── Step 1: Search ────────────────────────────────────────────
            search_context = ""
            results_by_query = {}

            search_available = SERPER_API_KEY or TAVILY_API_KEY
            if search_available and mode != "docs_simplifier" or (mode == "docs_simplifier" and not file_content):
                queries = build_search_queries(mode, topic, custom_instruction)
                yield f"data: {json.dumps({'type': 'status', 'message': 'Searching the web...', 'queries': queries})}\n\n"

                for q in queries:
                    results = web_search(q, num_results=6)
                    results_by_query[q] = results
                    yield f"data: {json.dumps({'type': 'search_done', 'query': q, 'count': len(results)})}\n\n"

                search_context = format_search_context(results_by_query)
                total_results = sum(len(v) for v in results_by_query.values())
                yield f"data: {json.dumps({'type': 'status', 'message': f'Found {total_results} sources. Generating report...'})}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'status', 'message': 'Generating report...'})}\n\n"

            # ── Step 2: Generate ──────────────────────────────────────────
            yield f"data: {json.dumps({'type': 'start', 'mode': mode, 'topic': topic})}\n\n"

            user_prompt = build_prompt(mode, topic, file_content, custom_instruction, search_context)

            stream = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=8000,
                temperature=0.7,
                stream=True
            )

            full_content = ""
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    text = chunk.choices[0].delta.content
                    full_content += text
                    yield f"data: {json.dumps({'type': 'chunk', 'text': text})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'full_content': full_content})}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )


@app.route('/api/export', methods=['POST'])
def export_report():
    data = request.json
    content = data.get('content', '')
    title = data.get('title', 'Research Report')
    export_format = data.get('format', 'md')

    title_clean = re.sub(r'[^\w\s-]', '', title)[:50].strip()

    if export_format == 'md':
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, mimetype='text/markdown', as_attachment=True,
                         download_name=f"{title_clean}.md")

    elif export_format == 'txt':
        plain = re.sub(r'#{1,6}\s', '', content)
        plain = re.sub(r'\*\*(.*?)\*\*', r'\1', plain)
        plain = re.sub(r'\*(.*?)\*', r'\1', plain)
        plain = re.sub(r'`(.*?)`', r'\1', plain)
        plain = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', plain)
        buffer = BytesIO(plain.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, mimetype='text/plain', as_attachment=True,
                         download_name=f"{title_clean}.txt")

    elif export_format == 'pdf':
        try:
            pdf_bytes = generate_pdf_from_markdown(content, title)
            buffer = BytesIO(pdf_bytes)
            buffer.seek(0)
            return send_file(buffer, mimetype='application/pdf', as_attachment=True,
                             download_name=f"{title_clean}.pdf")
        except Exception as e:
            return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500

    elif export_format == 'docx':
        try:
            docx_bytes = generate_docx_from_markdown(content, title)
            buffer = BytesIO(docx_bytes)
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{title_clean}.docx"
            )
        except Exception as e:
            return jsonify({"error": f"DOCX generation failed: {str(e)}"}), 500

    return jsonify({"error": "Invalid format"}), 400


if __name__ == '__main__':
    app.run(debug=True, port=5000)